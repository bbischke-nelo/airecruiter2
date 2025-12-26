#!/usr/bin/env python3
"""Test DOCX to PDF conversion.

Usage:
    PYTHONPATH=. ./venv/bin/python scripts/test_docx_conversion.py
    PYTHONPATH=. ./venv/bin/python scripts/test_docx_conversion.py --file /path/to/resume.docx
"""

import argparse
import io
import os


def create_sample_docx() -> bytes:
    """Create a sample DOCX for testing."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Add title
    title = doc.add_heading("John Doe", level=0)

    # Contact info
    doc.add_paragraph("john.doe@email.com | (555) 123-4567 | San Francisco, CA")

    # Experience section
    doc.add_heading("Experience", level=1)

    doc.add_heading("Senior Software Engineer", level=2)
    p = doc.add_paragraph()
    p.add_run("Acme Corp").bold = True
    p.add_run(" | 2020 - Present")
    doc.add_paragraph("- Led development of microservices architecture")
    doc.add_paragraph("- Managed team of 5 engineers")
    doc.add_paragraph("- Reduced deployment time by 80%")

    doc.add_heading("Software Engineer", level=2)
    p = doc.add_paragraph()
    p.add_run("StartupXYZ").bold = True
    p.add_run(" | 2017 - 2020")
    doc.add_paragraph("- Built REST APIs serving 10M+ requests/day")
    doc.add_paragraph("- Implemented CI/CD pipelines")

    # Education section
    doc.add_heading("Education", level=1)
    p = doc.add_paragraph()
    p.add_run("BS Computer Science").bold = True
    p.add_run(" - Stanford University, 2017")

    # Skills section
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, JavaScript, Go, AWS, Kubernetes, PostgreSQL, Redis")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def main(docx_path: str = None):
    from processor.utils.docx_to_pdf import convert_docx_to_pdf, is_docx

    if docx_path:
        print(f"Loading DOCX from: {docx_path}")
        with open(docx_path, "rb") as f:
            docx_content = f.read()
        filename = os.path.basename(docx_path)
    else:
        print("Creating sample DOCX...")
        docx_content = create_sample_docx()
        filename = "sample_resume.docx"

    print(f"DOCX size: {len(docx_content)} bytes")
    print(f"Is DOCX: {is_docx(filename)}")

    print("\nConverting to PDF...")
    pdf_content, new_filename, content_type = convert_docx_to_pdf(docx_content, filename)

    print(f"\nConversion successful!")
    print(f"  Original: {filename}")
    print(f"  New file: {new_filename}")
    print(f"  Content-Type: {content_type}")
    print(f"  PDF size: {len(pdf_content)} bytes")

    # Verify PDF header
    if pdf_content[:4] == b"%PDF":
        print("  PDF header: Valid")
    else:
        print(f"  PDF header: INVALID (got {pdf_content[:20]})")

    # Save to temp file for inspection
    output_path = f"/tmp/{new_filename}"
    with open(output_path, "wb") as f:
        f.write(pdf_content)
    print(f"\nSaved to: {output_path}")
    print("Open this file to verify the conversion quality.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", "-f", help="Path to existing DOCX file to convert")
    args = parser.parse_args()

    main(args.file)
