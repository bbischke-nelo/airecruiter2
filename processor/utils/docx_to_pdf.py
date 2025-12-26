"""Convert DOCX files to PDF for browser preview."""

import io
from typing import Tuple

import structlog

logger = structlog.get_logger()


def convert_docx_to_pdf(content: bytes, filename: str) -> Tuple[bytes, str, str]:
    """Convert a DOCX file to PDF.

    Uses python-docx to read the document and weasyprint to generate PDF.
    This allows browser-native preview without requiring LibreOffice.

    Args:
        content: DOCX file content as bytes
        filename: Original filename

    Returns:
        Tuple of (pdf_content, new_filename, content_type)

    Raises:
        ConversionError: If conversion fails
    """
    try:
        from docx import Document
        from weasyprint import HTML, CSS

        logger.info("Converting DOCX to PDF", filename=filename)

        # Read the DOCX
        doc = Document(io.BytesIO(content))

        # Convert to HTML
        html_content = _docx_to_html(doc)

        # Convert HTML to PDF with weasyprint
        html = HTML(string=html_content)
        css = CSS(string=_get_resume_css())
        pdf_bytes = html.write_pdf(stylesheets=[css])

        # Generate new filename
        new_filename = filename.rsplit(".", 1)[0] + ".pdf"

        logger.info(
            "DOCX converted to PDF",
            original_filename=filename,
            new_filename=new_filename,
            original_size=len(content),
            pdf_size=len(pdf_bytes),
        )

        return pdf_bytes, new_filename, "application/pdf"

    except ImportError as e:
        logger.error("Missing dependency for DOCX conversion", error=str(e))
        raise ConversionError(f"Missing dependency: {str(e)}") from e
    except Exception as e:
        logger.error("DOCX to PDF conversion failed", error=str(e), filename=filename)
        raise ConversionError(f"Conversion failed: {str(e)}") from e


def _docx_to_html(doc) -> str:
    """Convert a python-docx Document to HTML."""
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head><meta charset='utf-8'></head>",
        "<body>",
    ]

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        # Detect heading styles
        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name:
            html_parts.append(f"<h1>{_escape_html(para.text)}</h1>")
        elif "heading 2" in style_name:
            html_parts.append(f"<h2>{_escape_html(para.text)}</h2>")
        elif "heading 3" in style_name:
            html_parts.append(f"<h3>{_escape_html(para.text)}</h3>")
        elif "title" in style_name:
            html_parts.append(f"<h1 class='title'>{_escape_html(para.text)}</h1>")
        else:
            # Check for inline formatting
            html_parts.append(f"<p>{_format_paragraph(para)}</p>")

    # Extract tables
    for table in doc.tables:
        html_parts.append("<table>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_text = _escape_html(cell.text.strip())
                html_parts.append(f"<td>{cell_text}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    html_parts.extend(["</body>", "</html>"])

    return "\n".join(html_parts)


def _format_paragraph(para) -> str:
    """Format a paragraph with inline styles (bold, italic)."""
    parts = []
    for run in para.runs:
        text = _escape_html(run.text)
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    return "".join(parts)


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _get_resume_css() -> str:
    """CSS for rendering resumes nicely."""
    return """
        @page {
            size: letter;
            margin: 0.75in;
        }
        body {
            font-family: 'Helvetica Neue', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }
        h1 {
            font-size: 18pt;
            margin-bottom: 0.3em;
            color: #1a1a1a;
        }
        h1.title {
            font-size: 22pt;
            text-align: center;
            margin-bottom: 0.5em;
        }
        h2 {
            font-size: 14pt;
            margin-top: 1em;
            margin-bottom: 0.3em;
            color: #2a2a2a;
            border-bottom: 1px solid #ccc;
            padding-bottom: 0.2em;
        }
        h3 {
            font-size: 12pt;
            margin-top: 0.8em;
            margin-bottom: 0.2em;
        }
        p {
            margin: 0.3em 0;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 0.5em 0;
        }
        td {
            padding: 0.3em 0.5em;
            border: 1px solid #ddd;
            vertical-align: top;
        }
        strong {
            font-weight: 600;
        }
    """


def is_docx(filename: str, content_type: str = None) -> bool:
    """Check if a file is a DOCX document."""
    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return True

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in ("docx", "doc")


class ConversionError(Exception):
    """Raised when file conversion fails."""

    pass
