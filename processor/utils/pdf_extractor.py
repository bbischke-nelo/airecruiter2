"""Text extraction from PDF and DOCX files."""

import io
from typing import Optional

import structlog

logger = structlog.get_logger()


def extract_text_from_file(
    content: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> str:
    """Extract text content from a file.

    Args:
        content: File content as bytes
        filename: Original filename (for extension detection)
        content_type: MIME type (optional)

    Returns:
        Extracted text content
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Detect by content type or extension
    if content_type == "application/pdf" or ext == "pdf":
        return _extract_from_pdf(content)
    elif content_type in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ) or ext in ("doc", "docx"):
        return _extract_from_docx(content)
    elif content_type == "text/plain" or ext == "txt":
        return content.decode("utf-8", errors="ignore")
    else:
        logger.warning(
            "Unknown file type, attempting PDF extraction",
            filename=filename,
            content_type=content_type,
        )
        return _extract_from_pdf(content)


def _extract_from_pdf(content: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n\n".join(text_parts)

        logger.info(
            "PDF text extracted",
            pages=len(reader.pages),
            chars=len(text),
        )

        return text

    except Exception as e:
        logger.error("PDF extraction failed", error=str(e))
        raise ExtractionError(f"PDF extraction failed: {str(e)}") from e


def _extract_from_docx(content: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n\n".join(text_parts)

        logger.info(
            "DOCX text extracted",
            paragraphs=len(doc.paragraphs),
            chars=len(text),
        )

        return text

    except Exception as e:
        logger.error("DOCX extraction failed", error=str(e))
        raise ExtractionError(f"DOCX extraction failed: {str(e)}") from e


class ExtractionError(Exception):
    """Raised when text extraction fails."""

    pass
